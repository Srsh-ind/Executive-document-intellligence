import os
import pandas as pd
from pypdf import PdfReader
from docx import Document

from text_utils import clean_text


def read_pdf(path):
    reader = PdfReader(path)
    text = []

    for index, page in enumerate(reader.pages):
        page_text = page.extract_text()

        if page_text:
            text.append(f"--- Page {index + 1} ---\n{page_text}")

    return clean_text("\n\n".join(text))


def read_docx(path):
    doc = Document(path)
    text = []

    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text.append(" | ".join(cells))

    return clean_text("\n".join(text))


def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as file:
        return clean_text(file.read())


def read_csv(path):
    df = pd.read_csv(path, skipinitialspace=True, on_bad_lines="skip")
    return df.to_string(index=False)


def read_excel(path):
    df = pd.read_excel(path)
    return df.to_string(index=False)


def read_document(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return read_pdf(path)

    if ext == ".docx":
        return read_docx(path)

    if ext == ".txt":
        return read_txt(path)

    if ext == ".csv":
        return read_csv(path)

    if ext in [".xlsx", ".xls"]:
        return read_excel(path)

    raise ValueError(f"Unsupported file type: {ext}")